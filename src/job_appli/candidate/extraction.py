"""Deterministic resume text extraction from PDF (PyMuPDF) and DOCX (python-docx)."""

import io

MIN_EXTRACTABLE_CHARS = 100


class ExtractionTooSparse(Exception):
    pass


def extract_text(content: bytes, filename: str) -> str:
    name = filename.lower()
    if name.endswith(".pdf"):
        text = _extract_pdf(content)
    elif name.endswith(".docx"):
        text = _extract_docx(content)
    else:
        raise ValueError(f"unsupported resume format: {filename}")
    if len(text.strip()) < MIN_EXTRACTABLE_CHARS:
        raise ExtractionTooSparse(
            "extracted text is nearly empty; the file may be scanned or image-based"
        )
    return text


def _extract_pdf(content: bytes) -> str:
    import fitz

    with fitz.open(stream=content, filetype="pdf") as doc:
        return "\n".join(page.get_text() for page in doc)


def _extract_docx(content: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(content))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)
